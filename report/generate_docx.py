import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colors for professional theme (Indigo & Slate)
COLOR_PRIMARY = RGBColor(79, 70, 229)    # #4f46e5 (Indigo)
COLOR_SECONDARY = RGBColor(14, 165, 233)  # #0ea5e9 (Sky Blue)
COLOR_DARK = RGBColor(15, 23, 42)        # #0f172a (Dark Slate)
COLOR_MUTED = RGBColor(71, 85, 105)      # #475569 (Muted Slate)

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Sets cell margins (padding) in dxa (1/20 of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_text_formatting(run, font_name='Inter', size_pt=11, bold=False, italic=False, color=COLOR_DARK):
    """Applies clean typography to a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def create_schema_table(doc, entity_name, columns_data):
    """Creates a beautifully-styled schema table for the ERD section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r = p.add_run(f"Table Schema: {entity_name}")
    apply_text_formatting(r, font_name='Inter', size_pt=12, bold=True, color=COLOR_PRIMARY)
    
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    
    # Setup headers
    hdr_cells = table.rows[0].cells
    headers = ["Attribute Name", "Data Type", "Key", "Description"]
    widths = [Inches(1.8), Inches(1.2), Inches(0.8), Inches(2.7)]
    
    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i]
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "4F46E5") # Primary Indigo
        set_cell_margins(hdr_cells[i])
        run = hdr_cells[i].paragraphs[0].runs[0]
        apply_text_formatting(run, font_name='Inter', size_pt=10, bold=True, color=RGBColor(255, 255, 255))
        
    for row_data in columns_data:
        row_cells = table.add_row().cells
        for i in range(4):
            row_cells[i].width = widths[i]
            row_cells[i].text = row_data[i]
            set_cell_margins(row_cells[i])
            p_cell = row_cells[i].paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(2)
            run = p_cell.runs[0] if p_cell.runs else p_cell.add_run(row_data[i])
            apply_text_formatting(run, font_name='Inter', size_pt=9.5, bold=(i==0), color=COLOR_DARK)
            if i == 2 and ("PK" in row_data[i] or "FK" in row_data[i]):
                run.bold = True
                run.font.color.rgb = COLOR_SECONDARY

def build_report():
    print("Reading capstone_report.md...")
    with open("report/capstone_report.md", "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Custom styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Inter'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_DARK

    # Title & Subtitle Parsing
    title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
    subtitle_match = re.search(r'^## (.+)$', content, re.MULTILINE)
    
    title_text = title_match.group(1) if title_match else "Enterprise Retail Analytics Engine"
    subtitle_text = subtitle_match.group(1) if subtitle_match else "University Capstone Project Report"
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(8)
    r_title = p_title.add_run(title_text)
    apply_text_formatting(r_title, font_name='Inter', size_pt=26, bold=True, color=COLOR_PRIMARY)
    
    # Document Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(28)
    r_sub = p_sub.add_run(subtitle_text)
    apply_text_formatting(r_sub, font_name='Inter', size_pt=14, italic=True, color=COLOR_SECONDARY)
    
    # Course Meta
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(24)
    r_meta = p_meta.add_run("Course: Data Warehousing and Business Intelligence\nSubmission Date: June 12, 2026\nUniversity Capstone Project Submission\n\n\nSubmitted By:\n\n")
    apply_text_formatting(r_meta, font_name='Inter', size_pt=10.5, color=COLOR_MUTED)
    
    # Highlighted Rana Talha Majid
    r_lead = p_meta.add_run("★ RANA TALHA MAJID (23-CS-127) — Project Lead\n")
    apply_text_formatting(r_lead, font_name='Inter', size_pt=12, bold=True, color=COLOR_PRIMARY)
    r_lead_email = p_meta.add_run("Email: 23-cs-127@students.uettaxila.edu.pk\n\n")
    apply_text_formatting(r_lead_email, font_name='Inter', size_pt=9.5, italic=True, color=COLOR_MUTED)
    
    # Other group partners
    r_members = p_meta.add_run(
        "Muhammad Ibtasam Ali (23-CS-88)\n"
        "Email: 23-cs-88@students.uettaxila.edu.pk\n\n"
        "Ahmed Muneer (23-CS-91)\n"
        "Email: 23-cs-91@students.uettaxila.edu.pk\n"
    )
    apply_text_formatting(r_members, font_name='Inter', size_pt=10, color=COLOR_DARK)
    
    doc.add_page_break()

    # Split into lines
    lines = content.split('\n')
    
    # Keep track of active sections
    in_abstract = False
    in_table = False
    in_code = False
    in_mermaid = False
    table_rows = []
    
    # We will ignore original title/subtitle lines
    skip_lines_regex = re.compile(r'^#\s|^##\sUniversity\sCapstone')

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check for skip title/subtitle
        if skip_lines_regex.match(line):
            i += 1
            continue
            
        # Skip horizontal dividers
        if line == '---' or line == '***':
            i += 1
            continue
            
        # Skip "End of Report" or statistics footer lines
        if "*End of Report*" in line or "**Word Count**:" in line:
            i += 1
            continue
            
        # Code block check
        if line.startswith('```'):
            if in_code or in_mermaid:
                in_code = False
                in_mermaid = False
            else:
                if 'mermaid' in line:
                    in_mermaid = True
                    # Let's skip raw mermaid and render database tables instead
                    # ERD Tables Data
                    erd_fct_sales = [
                        ["sales_key", "varchar", "PK", "MD5 Hash Surrogate Key"],
                        ["customer_key", "varchar", "FK", "Reference to DIM_CUSTOMER"],
                        ["product_key", "varchar", "FK", "Reference to DIM_PRODUCT"],
                        ["date_key", "number", "FK", "Reference to DIM_DATE"],
                        ["order_id", "varchar", "DD", "Degenerate order natural ID"],
                        ["order_item_id", "varchar", "DD", "Degenerate item natural ID"],
                        ["quantity", "int", "", "Quantity ordered"],
                        ["unit_price", "number", "", "Discounted purchase price"],
                        ["revenue", "number", "", "Total revenue generated (quantity * price)"],
                        ["cogs", "number", "", "Cost of Goods Sold (quantity * cost)"],
                        ["gross_margin", "number", "", "Profit margin value (revenue - COGS)"],
                        ["gross_margin_pct", "number", "", "Margin rate (gross_margin / revenue * 100)"],
                        ["discount_amount", "number", "", "Value given away from standard retail"],
                        ["competitor_price_delta", "number", "", "Our price vs average competitor price"]
                    ]
                    erd_dim_customer = [
                        ["customer_key", "varchar", "PK", "MD5 Hash Surrogate Key"],
                        ["customer_id", "varchar", "", "Customer natural ID"],
                        ["full_name", "varchar", "", "Formatted Full Name"],
                        ["gender", "varchar", "", "Gender field (M / F / Non-Binary)"],
                        ["email", "varchar", "", "Email address"],
                        ["city", "varchar", "", "City location"],
                        ["country", "varchar", "", "Country weighted geographic distribution"],
                        ["registration_date", "date", "", "Account creation date"],
                        ["tenure_segment", "varchar", "", "Registration-based tenure group"],
                        ["activity_status", "varchar", "", "Recency status (Active / Inactive)"]
                    ]
                    erd_dim_product = [
                        ["product_key", "varchar", "PK", "MD5 Hash Surrogate Key"],
                        ["product_id", "varchar", "", "Product natural ID"],
                        ["product_name", "varchar", "", "Full item catalog name"],
                        ["category", "varchar", "", "Category classification (10 distinct categories)"],
                        ["retail_price", "number", "", "Standard retail unit price"],
                        ["base_cost", "number", "", "Product unit purchase cost"],
                        ["gross_margin_pct", "number", "", "Profitability margin profile"],
                        ["price_tier", "varchar", "", "Budget / Mid-Range / Premium / Luxury"],
                        ["competitive_position", "varchar", "", "Overpriced / Underpriced / Competitive"],
                        ["avg_competitor_price", "number", "", "Averaged scraped competitor price"]
                    ]
                    erd_dim_date = [
                        ["date_key", "number", "PK", "YYYYMMDD Integer Key"],
                        ["date_actual", "date", "", "Actual Calendar Date"],
                        ["year_number", "int", "", "Calendar Year"],
                        ["quarter_number", "int", "", "Calendar Quarter (1–4)"],
                        ["month_number", "int", "", "Calendar Month Number (1–12)"],
                        ["month_name", "varchar", "", "Month Word (e.g., January)"],
                        ["day_name", "varchar", "", "Day of the week (e.g., Monday)"],
                        ["is_weekend", "boolean", "", "True/False weekend flag"],
                        ["is_holiday_season", "boolean", "", "True/False holiday seasonal multiplier"]
                    ]
                    
                    create_schema_table(doc, "FCT_SALES (Fact Table)", erd_fct_sales)
                    create_schema_table(doc, "DIM_CUSTOMER (Dimension)", erd_dim_customer)
                    create_schema_table(doc, "DIM_PRODUCT (Dimension)", erd_dim_product)
                    create_schema_table(doc, "DIM_DATE (Dimension)", erd_dim_date)
                    
                    p_rel = doc.add_paragraph()
                    p_rel.paragraph_format.space_before = Pt(12)
                    p_rel.paragraph_format.space_after = Pt(4)
                    r_rel = p_rel.add_run("Schema Relationships:")
                    apply_text_formatting(r_rel, font_name='Inter', size_pt=11.5, bold=True, color=COLOR_PRIMARY)
                    
                    rel_items = [
                        "FCT_SALES.customer_key links many-to-one with DIM_CUSTOMER.customer_key",
                        "FCT_SALES.product_key links many-to-one with DIM_PRODUCT.product_key",
                        "FCT_SALES.date_key links many-to-one with DIM_DATE.date_key"
                    ]
                    for item in rel_items:
                        p_item = doc.add_paragraph(style='List Bullet')
                        p_item.paragraph_format.space_after = Pt(2)
                        run_item = p_item.add_run(item)
                        apply_text_formatting(run_item, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
                else:
                    in_code = True
                    # Plain text representation of code blocks (SQL, bash)
            i += 1
            continue
            
        if in_mermaid:
            # Skip lines in raw mermaid block
            i += 1
            continue
            
        if in_code:
            # Code block paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(4)
            set_cell_background(doc.add_table(rows=1, cols=1).rows[0].cells[0], "F1F5F9") # soft gray
            # Since document tables are simpler to handle, let's just write raw code block inside normal paragraph styled as monospaced
            p_code = doc.add_paragraph()
            p_code.paragraph_format.left_indent = Inches(0.4)
            p_code.paragraph_format.space_after = Pt(4)
            run = p_code.add_run(lines[i])
            apply_text_formatting(run, font_name='Consolas', size_pt=9.5, color=COLOR_MUTED)
            i += 1
            continue

        # Parsing table lines
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            i += 1
            continue
        elif in_table:
            # Table ended, render it
            in_table = False
            # Filter divider row (e.g. |---|---|)
            filtered_rows = [r for r in table_rows if not all(re.match(r'^[-:]+$', cell) for cell in r)]
            
            if filtered_rows:
                cols_count = len(filtered_rows[0])
                table = doc.add_table(rows=len(filtered_rows), cols=cols_count)
                table.autofit = True
                
                # Header row
                hdr_cells = table.rows[0].cells
                for col_idx, text in enumerate(filtered_rows[0]):
                    hdr_cells[col_idx].text = text
                    set_cell_background(hdr_cells[col_idx], "0F172A") # Dark slate header
                    set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=120, right=120)
                    run = hdr_cells[col_idx].paragraphs[0].runs[0]
                    apply_text_formatting(run, font_name='Inter', size_pt=9.5, bold=True, color=RGBColor(255, 255, 255))
                
                # Data rows
                for row_idx, row_data in enumerate(filtered_rows[1:]):
                    row_cells = table.rows[row_idx + 1].cells
                    # Zebra striping hex background
                    bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
                    
                    for col_idx, text in enumerate(row_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = text
                            set_cell_background(row_cells[col_idx], bg_color)
                            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
                            p_cell = row_cells[col_idx].paragraphs[0]
                            p_cell.paragraph_format.space_after = Pt(2)
                            if p_cell.runs:
                                apply_text_formatting(p_cell.runs[0], font_name='Inter', size_pt=9, color=COLOR_DARK)
                doc.add_paragraph().paragraph_format.space_after = Pt(6) # Spacing after table
            i += 1
            continue

        # Parsing Headings
        m_h1 = re.match(r'^##\s(\d+\.\s)?(.+)$', line)
        m_h2 = re.match(r'^###\s(\d+\.\d+\s)?(.+)$', line)
        
        if m_h1:
            title = m_h1.group(2)
            # Add section prefix if present
            if m_h1.group(1):
                title = f"{m_h1.group(1)}{title}"
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            apply_text_formatting(run, font_name='Inter', size_pt=15, bold=True, color=COLOR_PRIMARY)
            
        elif m_h2:
            title = m_h2.group(2)
            if m_h2.group(1):
                title = f"{m_h2.group(1)}{title}"
                
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title)
            apply_text_formatting(run, font_name='Inter', size_pt=12.5, bold=True, color=COLOR_SECONDARY)
            
        # List items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            
            cleaned_text = line[2:]
            
            # Format bold prefixes e.g. - **Data Silos**: text
            bold_prefix = re.match(r'^\*\*(.+?)\*\*:\s*(.*)', cleaned_text)
            if bold_prefix:
                run_b = p.add_run(bold_prefix.group(1) + ": ")
                apply_text_formatting(run_b, font_name='Inter', size_pt=10.5, bold=True, color=COLOR_DARK)
                run_t = p.add_run(bold_prefix.group(2))
                apply_text_formatting(run_t, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
            else:
                run = p.add_run(cleaned_text)
                apply_text_formatting(run, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
                
        # Numbered lists e.g. 1. text
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            
            cleaned_text = re.sub(r'^\d+\.\s', '', line)
            
            bold_prefix = re.match(r'^\*\*(.+?)\*\*:\s*(.*)', cleaned_text)
            if bold_prefix:
                run_b = p.add_run(bold_prefix.group(1) + ": ")
                apply_text_formatting(run_b, font_name='Inter', size_pt=10.5, bold=True, color=COLOR_DARK)
                run_t = p.add_run(bold_prefix.group(2))
                apply_text_formatting(run_t, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
            else:
                run = p.add_run(cleaned_text)
                apply_text_formatting(run, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
                
        # Normal paragraphs
        elif line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            # Format any simple bold markers inside paragraph
            # We split by ** and format alternating runs
            parts = re.split(r'\*\*([^*]+)\*\*', line)
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    run = p.add_run(part)
                    apply_text_formatting(run, font_name='Inter', size_pt=10.5, bold=True, color=COLOR_DARK)
                else:
                    # Clean inline links [link text](url) to just link text
                    clean_links = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', part)
                    run = p.add_run(clean_links)
                    apply_text_formatting(run, font_name='Inter', size_pt=10.5, color=COLOR_DARK)
                    
        i += 1

    out_path = "report/capstone_report.docx"
    print(f"Saving compiled docx to {out_path}...")
    doc.save(out_path)
    print("Document successfully compiled!")

if __name__ == "__main__":
    build_report()
